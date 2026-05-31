"""
Conversor Markdown — Núcleo de conversão v3
Sem dependências de GUI. Importável em:
  Desktop  : customtkinter / tkinter  (converter_app.py)
  Android  : Kivy, BeeWare Briefcase ou Chaquopy (Java → Python)

Para Android: importe este módulo e chame convert_file().
API estável — não muda entre versões de interface.
"""

import io
from pathlib import Path

SUPPORTED_EXTENSIONS = frozenset({
    ".html", ".htm", ".docx", ".pdf",
    ".csv", ".xlsx", ".xls", ".txt", ".rst", ".log",
})

EXT_LABEL = {
    ".pdf": "PDF",  ".docx": "Word",  ".html": "HTML",  ".htm": "HTML",
    ".xlsx": "Excel", ".xls": "Excel", ".csv": "CSV",
    ".txt": "Texto",  ".rst": "Texto", ".log": "Log",
}


# ── Detecção inteligente de encoding (fix #5 — acentos em Windows) ────────────
def _read_smart(path: Path) -> str:
    """
    Tenta UTF-8 → UTF-8 BOM → Windows-1252 → Latin-1.
    Resolve arquivos .txt criados no Windows (cp1252) sem quebrar acentos.
    """
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("latin-1")   # latin-1 nunca falha


# ── Conversores ───────────────────────────────────────────────────────────────

def convert_html(content: str) -> str:
    import html2text
    h = html2text.HTML2Text()
    h.ignore_links  = False
    h.ignore_images = False
    h.body_width    = 0
    return h.handle(content)


def convert_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    heading_map = {
        "Heading 1": "# ",  "Heading 2": "## ", "Heading 3": "### ",
        "Heading 4": "#### ", "Heading 5": "##### ", "Heading 6": "###### ",
        "Title": "# ", "Subtitle": "## ",
    }
    bullet_styles = {"List Bullet", "List Bullet 2", "List Bullet 3", "List Paragraph"}
    number_styles = {"List Number", "List Number 2", "List Number 3"}
    lines = []
    for para in doc.paragraphs:
        style  = para.style.name if para.style else "Normal"
        if style in bullet_styles:
            prefix = "- "
        elif style in number_styles:
            prefix = "1. "
        else:
            prefix = heading_map.get(style, "")
        parts = []
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                t = f"***{t}***"
            elif run.bold:
                t = f"**{t}**"
            elif run.italic:
                t = f"*{t}*"
            parts.append(t)
        line = prefix + "".join(parts)
        if line.strip():
            lines.append(line)
    for table in doc.tables:
        if not table.rows:
            continue
        # Deduplica células mescladas horizontalmente (python-docx repete o texto)
        def dedup_row(cells):
            seen, result = set(), []
            for i, c in enumerate(cells):
                t = c.text.strip()
                result.append(t if (t not in seen or t == "") else "")
                seen.add(t)
            return result
        cols = dedup_row(table.rows[0].cells)
        lines += [
            "",
            "| " + " | ".join(cols) + " |",
            "| " + " | ".join(["---"] * len(cols)) + " |",
        ]
        for row in table.rows[1:]:
            lines.append("| " + " | ".join(dedup_row(row.cells)) + " |")
        lines.append("")
    return "\n\n".join(l for l in lines if l.strip())


def convert_pdf(path: str) -> str:
    from pdfminer.high_level import extract_text
    text = extract_text(path)
    if not text.strip():
        return "_PDF sem texto extraível (pode ser imagem escaneada)._"
    return "\n".join(l.strip() for l in text.splitlines())


def convert_csv(path: str) -> str:
    import pandas as pd
    # Detecta encoding igual ao _read_smart — resolve CSV exportado pelo Excel (cp1252)
    raw = Path(path).read_bytes()
    text = None
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    df  = pd.read_csv(io.StringIO(text))
    sep = "| " + " | ".join(["---"] * len(df.columns)) + " |"
    rows = ["| " + " | ".join(str(v) for v in r) + " |" for _, r in df.iterrows()]
    return "\n".join(["| " + " | ".join(str(c) for c in df.columns) + " |", sep] + rows)


def convert_xlsx(path: str) -> str:
    import pandas as pd
    parts = []
    with pd.ExcelFile(path) as xl:
        for sheet in xl.sheet_names:
            df   = xl.parse(sheet)
            sep  = "| " + " | ".join(["---"] * len(df.columns)) + " |"
            rows = ["| " + " | ".join(str(v) for v in r) + " |" for _, r in df.iterrows()]
            parts.append(
                f"## {sheet}\n\n" +
                "\n".join(["| " + " | ".join(str(c) for c in df.columns) + " |", sep] + rows)
            )
    return "\n\n".join(parts)


# ── API pública ───────────────────────────────────────────────────────────────

def convert_file(input_path: str, output_path: str = None) -> tuple[str, str]:
    """
    Converte qualquer documento suportado para Markdown.
    Retorna (texto_markdown, caminho_salvo).

    Sem chamadas de GUI — compatível com Desktop e Android.
    """
    path = Path(input_path)
    ext  = path.suffix.lower()

    converters = {
        ".html": lambda: convert_html(_read_smart(path)),
        ".htm":  lambda: convert_html(_read_smart(path)),
        ".docx": lambda: convert_docx(str(path)),
        ".pdf":  lambda: convert_pdf(str(path)),
        ".csv":  lambda: convert_csv(str(path)),
        ".xlsx": lambda: convert_xlsx(str(path)),
        ".xls":  lambda: convert_xlsx(str(path)),
        ".txt":  lambda: _read_smart(path),
        ".rst":  lambda: _read_smart(path),
        ".log":  lambda: _read_smart(path),
    }

    if ext not in converters:
        raise ValueError(f"Formato '{ext}' não suportado.")

    markdown = converters[ext]()
    out = Path(output_path) if output_path else path.with_suffix(".md")
    out.write_text(markdown, encoding="utf-8")
    return markdown, str(out)
